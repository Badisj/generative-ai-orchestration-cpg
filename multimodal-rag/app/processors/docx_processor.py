# FILE: app/processors/docx_processor.py
"""
Word document processor using python-docx.
Extracts text with structure preservation.
"""

from pathlib import Path
from typing import List
import logging

from app.processors.base import BaseProcessor, Document

logger = logging.getLogger(__name__)


class DocxProcessor(BaseProcessor):
    """
    Process Word documents (.docx, .doc).
    
    Features:
    - Extracts paragraphs with formatting context
    - Handles tables
    - Preserves headings and structure
    """
    
    supported_extensions = [".docx", ".doc"]
    
    def can_process(self, file_path: Path) -> bool:
        """Check if file is a Word document."""
        return file_path.suffix.lower() in self.supported_extensions
    
    def process(self, file_path: Path) -> List[Document]:
        """
        Extract text from Word document.
        
        Parameters
        ----------
        file_path : Path
            Path to Word file.
            
        Returns
        -------
        List[Document]
            Single document with all extracted content.
        """
        from docx import Document as DocxDocument
        from docx.opc.exceptions import PackageNotFoundError
        
        file_path = Path(file_path)
        
        # Handle .doc files (older format)
        if file_path.suffix.lower() == ".doc":
            return self._process_doc(file_path)
        
        try:
            doc = DocxDocument(str(file_path))
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Check if it's a heading
                    if para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading ", "")
                        try:
                            level = int(level)
                            prefix = "#" * level + " "
                        except ValueError:
                            prefix = "## "
                        content_parts.append(f"{prefix}{text}")
                    else:
                        content_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_md = self._table_to_markdown(table)
                if table_md:
                    content_parts.append(table_md)
            
            content = "\n\n".join(content_parts)
            
            if content:
                document = self._create_document(
                    content=content,
                    source_file=file_path.name,
                    file_type="docx",
                    extraction_method="python-docx",
                )
                logger.info(f"Extracted content from {file_path.name}")
                return [document]
            
        except PackageNotFoundError:
            logger.error(f"Invalid or corrupted Word file: {file_path}")
        except Exception as e:
            logger.error(f"Error processing Word file {file_path}: {e}")
        
        return []
    
    def _process_doc(self, file_path: Path) -> List[Document]:
        """
        Process older .doc format using antiword or fallback.
        
        Parameters
        ----------
        file_path : Path
            Path to .doc file.
            
        Returns
        -------
        List[Document]
            Extracted documents.
        """
        try:
            # Try using docx2txt as fallback
            import docx2txt
            text = docx2txt.process(str(file_path))
            
            if text and text.strip():
                document = self._create_document(
                    content=text.strip(),
                    source_file=file_path.name,
                    file_type="doc",
                    extraction_method="docx2txt",
                )
                return [document]
                
        except ImportError:
            logger.warning("docx2txt not installed for .doc support")
        except Exception as e:
            logger.warning(f"Could not process .doc file: {e}")
        
        return []
    
    def _table_to_markdown(self, table) -> str:
        """
        Convert Word table to markdown format.
        
        Parameters
        ----------
        table : docx.table.Table
            Word table object.
            
        Returns
        -------
        str
            Markdown formatted table.
        """
        rows = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            
            # Add header separator after first row
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                rows.append(separator)
        
        return "\n".join(rows) if rows else ""
